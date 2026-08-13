"""
framework.report
================
Builds the assessment report from session data. The report is assembled
as structured sections, then rendered to Markdown, styled HTML (printable
to PDF via the browser), and JSON.
"""

import html
import json
from datetime import date

from . import reference as ref
from . import scoring


# ---------------------------------------------------------------------------
# Assemble structured report data from app state (a plain dict)
# ---------------------------------------------------------------------------

def assemble(state: dict) -> dict:
    classification = ref.CLASSIFICATION_MATRIX[(state["mechanism"], state["orientation"])]
    buckets = scoring.group_by_bucket(state["selected_indicators"])
    checks = scoring.run_integrity_checks(
        state["selected_indicators"], state["uncertainty"], state.get("pathway_stage_names", [])
    )
    used_refs = sorted({e.get("citation") for e in state["selected_indicators"].values()
                        if e.get("citation")} |
                       {"weiss1995", "mayne2008", "kellogg2004", "mclaughlin2015",
                        "iso2015", "sahili2024", "lilin2025", "roomi2021", "weidema2018"})
    return {
        "meta": {
            "startup_name": state["startup_name"],
            "startup_desc": state["startup_desc"],
            "sector": state.get("sector", ""),
            "stage": state["stage"],
            "generated": date.today().isoformat(),
            "assessment_version": state.get("assessment_version", 1),
            "next_review_milestone": state.get("next_review_milestone", ""),
            "review_notes": state.get("review_notes", ""),
        },
        "classification": {
            "mechanism": state["mechanism"],
            "orientation": state["orientation"],
            "is_hybrid": state.get("is_hybrid", False),
            "secondary_mechanism": state.get("secondary_mechanism"),
            "label": classification["label"],
            "blurb": classification["blurb"],
            "unit_of_analysis": classification["unit_of_analysis"],
        },
        "pathway": state.get("pathway", {}),           # {track: [{stage, description, assumption, evidence}]}
        "weakest_links": state.get("weakest_links", ""),
        "buckets": {b: [(n, e) for n, e in items] for b, items in buckets.items()},
        "uncertainty": state["uncertainty"],
        "integrity_checks": checks,
        "references": {k: ref.REFERENCES[k] for k in used_refs if k in ref.REFERENCES},
        "stage_note": ref.STAGE_GUIDANCE[state["stage"]]["note"],
    }


# ---------------------------------------------------------------------------
# Markdown rendering
# ---------------------------------------------------------------------------

def to_markdown(r: dict) -> str:
    m, c = r["meta"], r["classification"]
    L = []
    L.append(f"# Environmental Impact Assessment — {m['startup_name']}")
    L.append(f"*Assessment version {m['assessment_version']} · Generated {m['generated']} · "
             f"Development stage: {m['stage']}*\n")
    if m["startup_desc"]:
        L.append(f"> {m['startup_desc']}\n")
    if m["sector"]:
        L.append(f"**Sector:** {m['sector']}\n")

    # Executive summary
    core = r["buckets"].get(scoring.BUCKET_CORE, [])
    L.append("## Executive Summary")
    L.append(f"- **Classification:** {c['label']}")
    L.append(f"- **Primary unit of analysis:** {c['unit_of_analysis']}")
    L.append(f"- **Core indicator set:** {len(core)} indicator(s)"
             + (": " + "; ".join(n for n, _ in core) if core else ""))
    lvl_counts = {}
    for u in r["uncertainty"].values():
        lvl_counts[u.get("level", "unlabelled")] = lvl_counts.get(u.get("level", "unlabelled"), 0) + 1
    if lvl_counts:
        L.append("- **Evidence base:** " + ", ".join(f"{v} {k.lower()}" for k, v in lvl_counts.items()))
    warn = [msg for sev, msg in r["integrity_checks"] if sev == "warning"]
    L.append(f"- **Integrity checks:** {len(warn)} warning(s)"
             + ("" if not warn else " — see Section 5"))
    if m["next_review_milestone"]:
        L.append(f"- **Next scheduled review:** {m['next_review_milestone']}")
    L.append("")

    # Module 1
    L.append("## 1. Startup Profiling & Type Classification")
    L.append(f"- **Impact mechanism:** {c['mechanism']}"
             + (f" (secondary track: {c['secondary_mechanism']})" if c["is_hybrid"] else ""))
    L.append(f"- **Value proposition orientation:** {c['orientation']}")
    L.append(f"- **Resulting type:** {c['label']}")
    L.append(f"- {c['blurb']}\n")
    L.append(f"*Stage guidance ({m['stage']}):* {r['stage_note']}\n")

    # Module 2
    L.append("## 2. Impact Pathway (Theory of Change)")
    for track, stages in r["pathway"].items():
        L.append(f"### {track}")
        for s in stages:
            L.append(f"- **{s['stage']}**")
            L.append(f"  - Description: {s.get('description') or '_(not filled in)_'}")
            L.append(f"  - Underlying assumption: {s.get('assumption') or '_(not stated)_'}")
            L.append(f"  - Evidence strength: {s.get('evidence', 'Not rated')}")
    if r["weakest_links"]:
        L.append(f"\n**Weakest link(s) in the chain:** {r['weakest_links']}")
    L.append("")

    # Module 3
    L.append("## 3. Indicator Set (Feasibility–Relevance Scoring)")
    for b in scoring.BUCKET_ORDER:
        items = r["buckets"].get(b, [])
        if not items:
            continue
        L.append(f"### {b}")
        for n, e in items:
            details = [f"Feasibility: {e['feasibility']}", f"Relevance: {e['relevance']}"]
            if e.get("unit"):
                details.append(f"Unit: {e['unit']}")
            if e.get("current_value"):
                details.append(f"Current value: {e['current_value']}")
            if e.get("target"):
                details.append(f"Target: {e['target']}")
            if e.get("frequency"):
                details.append(f"Measured: {e['frequency']}")
            if e.get("data_source"):
                details.append(f"Data source: {e['data_source']}")
            if e.get("pathway_link"):
                details.append(f"Evidences pathway stage: {e['pathway_link']}")
            L.append(f"- **{n}** ({e['category']}) — " + " · ".join(details))
    L.append("")

    # Module 4
    L.append("## 4. Uncertainty Acknowledgement")
    if not r["uncertainty"]:
        L.append("_No claims labelled yet._")
    for n, u in r["uncertainty"].items():
        L.append(f"### {n}")
        lvl = u.get("level", "unlabelled")
        L.append(f"- **Evidential confidence:** {lvl}"
                 + (f" — {ref.CONFIDENCE_LEVELS[lvl]}" if lvl in ref.CONFIDENCE_LEVELS else ""))
        if u.get("claim"):
            L.append(f"- **Claim as stated:** {u['claim']}")
        L.append(f"- **Assumptions & data sources:** {u.get('assumptions') or '_(not filled in)_'}")
        L.append(f"- **Conditions under which the estimate could differ significantly:** "
                 f"{u.get('conditions') or '_(not filled in)_'}")
    L.append("")

    # Integrity checks
    L.append("## 5. Integrity Checks")
    sym = {"warning": "⚠️", "info": "ℹ️", "ok": "✅"}
    for sev, msg in r["integrity_checks"]:
        L.append(f"- {sym[sev]} {msg}")
    L.append("")

    # Review plan
    L.append("## 6. Review & Updating Plan")
    L.append("This assessment is designed for updating, not for completion. It should "
             "be revisited at defined milestones rather than filed as a one-time exercise.")
    L.append(f"- **Next review milestone:** {m['next_review_milestone'] or '_(not set)_'}")
    if m["review_notes"]:
        L.append(f"- **Notes:** {m['review_notes']}")
    L.append("")

    # Methodology + references
    L.append("## Methodology")
    L.append(ref.METHODOLOGY_NOTE)
    L.append("")
    L.append("## References")
    for k in sorted(r["references"]):
        L.append(f"- {r['references'][k]}")
    return "\n".join(L)


# ---------------------------------------------------------------------------
# HTML rendering (self-contained, printable to PDF from the browser)
# ---------------------------------------------------------------------------

_CSS = """
body { font-family: Georgia, 'Times New Roman', serif; max-width: 820px;
       margin: 2.5rem auto; padding: 0 1.5rem; color: #1a2b23; line-height: 1.55; }
h1 { font-size: 1.7rem; border-bottom: 3px solid #2e6b4f; padding-bottom: .4rem; }
h2 { font-size: 1.25rem; color: #2e6b4f; margin-top: 2rem;
     border-bottom: 1px solid #cfe0d6; padding-bottom: .2rem; }
h3 { font-size: 1.02rem; margin-bottom: .2rem; }
.meta { color: #5a6b62; font-style: italic; }
blockquote { border-left: 4px solid #2e6b4f; margin: 1rem 0; padding: .3rem 1rem;
             background: #f2f7f4; }
ul { padding-left: 1.3rem; }
li { margin: .25rem 0; }
.badge { display: inline-block; padding: .1rem .55rem; border-radius: 1rem;
         font-size: .8rem; font-family: Arial, sans-serif; color: #fff; }
.b-Measured { background: #2e6b4f; } .b-Modelled { background: #b8860b; }
.b-Projected { background: #8a5a9e; }
.warn { color: #a33; } .ok { color: #2e6b4f; } .info { color: #345d8a; }
footer { margin-top: 3rem; font-size: .85rem; color: #5a6b62;
         border-top: 1px solid #cfe0d6; padding-top: .8rem; }
@media print { body { margin: 0 auto; } }
"""


def to_html(r: dict) -> str:
    """Render the report to a self-contained HTML document."""
    m, c = r["meta"], r["classification"]
    e = html.escape

    def li(items):
        return "<ul>" + "".join(f"<li>{i}</li>" for i in items) + "</ul>"

    parts = [f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
             f"<title>Environmental Impact Assessment — {e(m['startup_name'])}</title>"
             f"<style>{_CSS}</style></head><body>"]
    parts.append(f"<h1>Environmental Impact Assessment — {e(m['startup_name'])}</h1>")
    parts.append(f"<p class='meta'>Assessment version {m['assessment_version']} · "
                 f"Generated {m['generated']} · Development stage: {e(m['stage'])}"
                 + (f" · Sector: {e(m['sector'])}" if m["sector"] else "") + "</p>")
    if m["startup_desc"]:
        parts.append(f"<blockquote>{e(m['startup_desc'])}</blockquote>")

    core = r["buckets"].get(scoring.BUCKET_CORE, [])
    warn_n = sum(1 for s, _ in r["integrity_checks"] if s == "warning")
    lvl_counts = {}
    for u in r["uncertainty"].values():
        lvl_counts[u.get("level", "unlabelled")] = lvl_counts.get(u.get("level", "unlabelled"), 0) + 1
    parts.append("<h2>Executive Summary</h2>")
    parts.append(li([
        f"<b>Classification:</b> {e(c['label'])}",
        f"<b>Primary unit of analysis:</b> {e(c['unit_of_analysis'])}",
        f"<b>Core indicator set:</b> {len(core)} indicator(s)"
        + (": " + "; ".join(e(n) for n, _ in core) if core else ""),
        "<b>Evidence base:</b> " + ", ".join(f"{v} {e(k.lower())}" for k, v in lvl_counts.items())
        if lvl_counts else "<b>Evidence base:</b> not yet labelled",
        f"<b>Integrity checks:</b> {warn_n} warning(s)",
        f"<b>Next scheduled review:</b> {e(m['next_review_milestone']) or '(not set)'}",
    ]))

    parts.append("<h2>1. Startup Profiling &amp; Type Classification</h2>")
    parts.append(li([
        f"<b>Impact mechanism:</b> {e(c['mechanism'])}"
        + (f" (secondary track: {e(c['secondary_mechanism'] or '')})" if c["is_hybrid"] else ""),
        f"<b>Value proposition orientation:</b> {e(c['orientation'])}",
        f"<b>Resulting type:</b> {e(c['label'])} — {e(c['blurb'])}",
    ]))
    parts.append(f"<p class='meta'>Stage guidance ({e(m['stage'])}): {e(r['stage_note'])}</p>")

    parts.append("<h2>2. Impact Pathway (Theory of Change)</h2>")
    for track, stages in r["pathway"].items():
        parts.append(f"<h3>{e(track)}</h3>")
        rows = []
        for s in stages:
            rows.append(f"<b>{e(s['stage'])}</b><br>"
                        f"Description: {e(s.get('description') or '(not filled in)')}<br>"
                        f"Assumption: {e(s.get('assumption') or '(not stated)')}<br>"
                        f"Evidence strength: {e(s.get('evidence', 'Not rated'))}")
        parts.append(li(rows))
    if r["weakest_links"]:
        parts.append(f"<p><b>Weakest link(s):</b> {e(r['weakest_links'])}</p>")

    parts.append("<h2>3. Indicator Set (Feasibility–Relevance Scoring)</h2>")
    for b in scoring.BUCKET_ORDER:
        items = r["buckets"].get(b, [])
        if not items:
            continue
        parts.append(f"<h3>{e(b)}</h3>")
        rows = []
        for n, en in items:
            det = [f"Feasibility: {en['feasibility']}", f"Relevance: {en['relevance']}"]
            for field, lbl in [("unit", "Unit"), ("current_value", "Current value"),
                               ("target", "Target"), ("frequency", "Measured"),
                               ("data_source", "Data source"), ("pathway_link", "Evidences stage")]:
                if en.get(field):
                    det.append(f"{lbl}: {e(str(en[field]))}")
            rows.append(f"<b>{e(n)}</b> <i>({e(en['category'])})</i> — " + " · ".join(det))
        parts.append(li(rows))

    parts.append("<h2>4. Uncertainty Acknowledgement</h2>")
    for n, u in r["uncertainty"].items():
        lvl = u.get("level", "unlabelled")
        badge = f"<span class='badge b-{lvl}'>{e(lvl)}</span>" if lvl in ("Measured", "Modelled", "Projected") else e(lvl)
        parts.append(f"<h3>{e(n)} {badge}</h3>")
        rows = []
        if u.get("claim"):
            rows.append(f"<b>Claim as stated:</b> {e(u['claim'])}")
        rows.append(f"<b>Assumptions &amp; data sources:</b> {e(u.get('assumptions') or '(not filled in)')}")
        rows.append(f"<b>Conditions that could change the estimate:</b> {e(u.get('conditions') or '(not filled in)')}")
        parts.append(li(rows))

    parts.append("<h2>5. Integrity Checks</h2>")
    cls = {"warning": "warn", "ok": "ok", "info": "info"}
    sym = {"warning": "⚠", "ok": "✔", "info": "ℹ"}
    parts.append(li([f"<span class='{cls[s]}'>{sym[s]} {e(msg)}</span>"
                     for s, msg in r["integrity_checks"]]))

    parts.append("<h2>6. Review &amp; Updating Plan</h2>")
    parts.append("<p>This assessment is designed for updating, not for completion; it should be "
                 "revisited at defined milestones rather than filed as a one-time exercise.</p>")
    parts.append(li([f"<b>Next review milestone:</b> {e(m['next_review_milestone']) or '(not set)'}"]
                    + ([f"<b>Notes:</b> {e(m['review_notes'])}"] if m["review_notes"] else [])))

    parts.append("<h2>Methodology</h2>")
    parts.append(f"<p>{e(ref.METHODOLOGY_NOTE)}</p>")
    parts.append("<h2>References</h2>")
    parts.append(li([e(v) for k, v in sorted(r["references"].items())]))
    parts.append("<footer>Generated with the four-module Environmental Impact Assessment "
                 "tool, based on “A Proposed Framework for Environmental Impact Assessment "
                 "in Early-Stage Startups”. Screening-level instrument — does not replace "
                 "full LCA or GHG Protocol reporting.</footer>")
    parts.append("</body></html>")
    return "".join(parts)


# ---------------------------------------------------------------------------
# Word (.docx) rendering — the primary human-readable deliverable.
# Built with python-docx: real heading styles, a table of contents field,
# styled indicator tables, colour-coded confidence badges, and a page-numbered
# footer. Returns raw bytes, ready for a Streamlit download_button.
# ---------------------------------------------------------------------------

_DOCX_GREEN = "2E6B4F"
_DOCX_GOLD = "B8860B"
_DOCX_PURPLE = "8A5A9E"
_DOCX_RED = "AA3333"
_DOCX_BLUE = "34608A"
_DOCX_GREY = "5A6B62"
_DOCX_LIGHT = "EAF2ED"

_CONFIDENCE_COLOR = {"Measured": _DOCX_GREEN, "Modelled": _DOCX_GOLD, "Projected": _DOCX_PURPLE}
_CHECK_COLOR = {"warning": _DOCX_RED, "info": _DOCX_BLUE, "ok": _DOCX_GREEN}
_CHECK_SYMBOL = {"warning": "\u26A0", "info": "\u2139", "ok": "\u2713"}


def _docx_set_base_style(document):
    """Give the document a cleaner default look than vanilla python-docx."""
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    title = document.styles["Title"]
    title.font.color.rgb = None  # keep default; set per-run below for control
    for lvl, size in ((1, 16), (2, 13), (3, 11.5)):
        h = document.styles[f"Heading {lvl}"]
        h.font.size = Pt(size)
        h.font.bold = True


def _docx_hex_color(run, hex_str):
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor.from_string(hex_str)


def _docx_add_toc(document):
    """Insert a real Word TOC field. Word populates it on open when the user
    right-clicks and chooses 'Update Field' (or opens with fields set to
    auto-update) — this is standard behaviour for field-based TOCs generated
    outside Word itself."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose \u201cUpdate Field\u201d to generate the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)  # sibling of the field-char markers, not nested inside one
    r.append(fld_end)


def _docx_add_page_numbers(document):
    """Footer: 'Page X of Y' via PAGE/NUMPAGES fields, plus a source line."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    def field(name):
        el = OxmlElement("w:fldSimple")
        el.set(qn("w:instr"), name)
        return el

    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = None
    p._p.append(field("PAGE"))
    run2 = p.add_run(" of ")
    run2.font.size = Pt(8.5)
    p._p.append(field("NUMPAGES"))


def _docx_bullets(document, items, style="List Bullet"):
    for item in items:
        document.add_paragraph(item, style=style)


def _docx_badge(paragraph, level):
    run = paragraph.add_run(f"  [{level.upper()}]")
    run.bold = True
    _docx_hex_color(run, _CONFIDENCE_COLOR.get(level, _DOCX_GREY))


def _docx_indicator_table(document, items):
    """items: list of (name, entry_dict)."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    headers = ["Indicator", "Category", "Feas.", "Rel.", "Details"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)

    for name, e in items:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = e.get("category", "")
        row[2].text = e.get("feasibility", "")
        row[3].text = e.get("relevance", "")
        details = []
        if e.get("unit"):
            details.append(f"Unit: {e['unit']}")
        if e.get("current_value"):
            details.append(f"Current: {e['current_value']}")
        if e.get("target"):
            details.append(f"Target: {e['target']}")
        if e.get("frequency"):
            details.append(f"Freq.: {e['frequency']}")
        if e.get("data_source"):
            details.append(f"Source: {e['data_source']}")
        if e.get("pathway_link"):
            details.append(f"Evidences: {e['pathway_link']}")
        row[4].text = "; ".join(details) if details else "\u2014"
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    # zebra-stripe body rows for readability
    for idx, row in enumerate(table.rows[1:], start=1):
        if idx % 2 == 0:
            for cell in row.cells:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), _DOCX_LIGHT)
                cell._tc.get_or_add_tcPr().append(shd)
    document.add_paragraph()  # spacing after table


def to_docx(r: dict) -> bytes:
    """Render the report as a Word document. Returns bytes suitable for a
    Streamlit download_button (mimetype application/vnd.openxmlformats-
    officedocument.wordprocessingml.document)."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    m, c = r["meta"], r["classification"]
    document = Document()
    _docx_set_base_style(document)

    # --- Title & meta -----------------------------------------------------
    title = document.add_heading(f"Environmental Impact Assessment", level=0)
    title.runs[0].font.color.rgb = RGBColor.from_string(_DOCX_GREEN)
    sub = document.add_heading(m["startup_name"] or "Untitled startup", level=1)
    sub.runs[0].font.color.rgb = RGBColor.from_string("333333")

    meta_p = document.add_paragraph()
    meta_run = meta_p.add_run(
        f"Assessment version {m['assessment_version']}  \u00b7  Generated {m['generated']}  \u00b7  "
        f"Development stage: {m['stage']}" + (f"  \u00b7  Sector: {m['sector']}" if m["sector"] else "")
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor.from_string(_DOCX_GREY)

    if m["startup_desc"]:
        quote = document.add_paragraph()
        quote.paragraph_format.left_indent = Pt(18)
        qrun = quote.add_run(m["startup_desc"])
        qrun.italic = True

    document.add_paragraph()
    document.add_heading("Table of Contents", level=1)
    _docx_add_toc(document)
    document.add_page_break()

    # --- Executive summary --------------------------------------------------
    document.add_heading("Executive Summary", level=1)
    core = r["buckets"].get(scoring.BUCKET_CORE, [])
    lvl_counts = {}
    for u in r["uncertainty"].values():
        lvl_counts[u.get("level", "unlabelled")] = lvl_counts.get(u.get("level", "unlabelled"), 0) + 1
    warn_n = sum(1 for sev, _ in r["integrity_checks"] if sev == "warning")

    summary_items = [
        f"Classification: {c['label']}",
        f"Primary unit of analysis: {c['unit_of_analysis']}",
        f"Core indicator set: {len(core)} indicator(s)"
        + (" \u2014 " + "; ".join(n for n, _ in core) if core else ""),
        "Evidence base: " + (", ".join(f"{v} {k.lower()}" for k, v in lvl_counts.items())
                             if lvl_counts else "not yet labelled"),
        f"Integrity checks: {warn_n} warning(s)" + (" \u2014 see Section 5" if warn_n else ""),
        f"Next scheduled review: {m['next_review_milestone'] or '(not set)'}",
    ]
    _docx_bullets(document, summary_items)

    # --- Module 1 ------------------------------------------------------------
    document.add_heading("1. Startup Profiling & Type Classification", level=1)
    mech_line = f"Impact mechanism: {c['mechanism']}"
    if c["is_hybrid"]:
        mech_line += f" (secondary track: {c['secondary_mechanism']})"
    _docx_bullets(document, [
        mech_line,
        f"Value proposition orientation: {c['orientation']}",
        f"Resulting type: {c['label']}",
    ])
    p = document.add_paragraph(c["blurb"])
    p.paragraph_format.left_indent = Pt(10)
    stage_p = document.add_paragraph()
    stage_run = stage_p.add_run(f"Stage guidance ({m['stage']}): {r['stage_note']}")
    stage_run.italic = True
    stage_run.font.size = Pt(9.5)
    stage_run.font.color.rgb = RGBColor.from_string(_DOCX_GREY)

    # --- Module 2 --------------------------------------------------------
    document.add_heading("2. Impact Pathway (Theory of Change)", level=1)
    for track, stages in r["pathway"].items():
        document.add_heading(track, level=2)
        for s in stages:
            document.add_heading(s["stage"], level=3)
            _docx_bullets(document, [
                f"Description: {s.get('description') or '(not filled in)'}",
                f"Underlying assumption: {s.get('assumption') or '(not stated)'}",
                f"Evidence strength: {s.get('evidence', 'Not rated')}",
            ])
    if r["weakest_links"]:
        wl = document.add_paragraph()
        wl_run = wl.add_run(f"Weakest link(s) in the chain: {r['weakest_links']}")
        wl_run.bold = True
        _docx_hex_color(wl_run, _DOCX_RED)

    # --- Module 3 --------------------------------------------------------
    document.add_heading("3. Indicator Set (Feasibility\u2013Relevance Scoring)", level=1)
    any_indicators = False
    for b in scoring.BUCKET_ORDER:
        items = r["buckets"].get(b, [])
        if not items:
            continue
        any_indicators = True
        document.add_heading(b, level=2)
        _docx_indicator_table(document, items)
    if not any_indicators:
        document.add_paragraph("No indicators selected.")

    # --- Module 4 --------------------------------------------------------
    document.add_heading("4. Uncertainty Acknowledgement", level=1)
    if not r["uncertainty"]:
        document.add_paragraph("No claims labelled yet.")
    for name, u in r["uncertainty"].items():
        h = document.add_heading(level=3)
        h.add_run(name)
        lvl = u.get("level", "unlabelled")
        _docx_badge(h, lvl)
        if u.get("claim"):
            cp = document.add_paragraph()
            cp.add_run("Claim as stated: ").bold = True
            cp.add_run(u["claim"])
        ap = document.add_paragraph()
        ap.add_run("Assumptions & data sources: ").bold = True
        ap.add_run(u.get("assumptions") or "(not filled in)")
        cnp = document.add_paragraph()
        cnp.add_run("Conditions that could change the estimate: ").bold = True
        cnp.add_run(u.get("conditions") or "(not filled in)")

    # --- Section 5: Integrity checks --------------------------------------
    document.add_heading("5. Integrity Checks", level=1)
    for sev, msg in r["integrity_checks"]:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(f"{_CHECK_SYMBOL[sev]}  {msg}")
        _docx_hex_color(run, _CHECK_COLOR[sev])

    # --- Section 6: Review plan --------------------------------------------
    document.add_heading("6. Review & Updating Plan", level=1)
    document.add_paragraph(
        "This assessment is designed for updating, not for completion: it should "
        "be revisited at defined milestones rather than filed as a one-time exercise."
    )
    review_items = [f"Next review milestone: {m['next_review_milestone'] or '(not set)'}"]
    if m["review_notes"]:
        review_items.append(f"Notes: {m['review_notes']}")
    _docx_bullets(document, review_items)

    # --- Methodology & references -----------------------------------------
    document.add_heading("Methodology", level=1)
    document.add_paragraph(ref.METHODOLOGY_NOTE)
    document.add_heading("References", level=1)
    for k in sorted(r["references"]):
        document.add_paragraph(r["references"][k], style="List Bullet")

    footer_p = document.add_paragraph()
    footer_run = footer_p.add_run(
        "Generated with the four-module Environmental Impact Assessment tool, based on "
        "\u201cA Proposed Framework for Environmental Impact Assessment in Early-Stage "
        "Startups\u201d. Screening-level instrument \u2014 does not replace full LCA or "
        "GHG Protocol reporting."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor.from_string(_DOCX_GREY)

    _docx_add_page_numbers(document)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# JSON export / import (state persistence for milestone-based updating)
# ---------------------------------------------------------------------------

PERSISTED_KEYS = [
    "startup_name", "startup_desc", "sector", "stage",
    "mechanism", "orientation", "is_hybrid", "secondary_mechanism",
    "pathway", "weakest_links",
    "selected_indicators", "custom_indicators", "uncertainty",
    "assessment_version", "next_review_milestone", "review_notes",
]


def export_state(state: dict) -> str:
    payload = {
        "format": "startup-env-impact-assessment",
        "format_version": 2,
        "exported": date.today().isoformat(),
        "state": {k: state.get(k) for k in PERSISTED_KEYS},
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def import_state(raw: str) -> dict:
    payload = json.loads(raw)
    if payload.get("format") != "startup-env-impact-assessment":
        raise ValueError("Not a recognised assessment file.")
    state = payload["state"]
    # bump the version on re-import: this is a new review cycle
    state["assessment_version"] = int(state.get("assessment_version", 1)) + 1
    return state
